"""
Word 文档解析器
支持：.docx 格式（.doc 需要转换）
"""

import logging
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class DOCXExtractionResult:
    """Word 文档解析结果"""
    file_path: str
    title: str = ""
    paragraphs: list[str] = field(default_factory=list)
    tables: list[str] = field(default_factory=list)  # Markdown 格式
    full_text: str = ""
    metadata: dict = field(default_factory=dict)


class DOCXExtractor:
    """
    Word 文档解析器
    
    优点：原生支持标题层级、表格、结构化内容
    缺点：不支持 .doc（需要 LibreOffice 转换）
    """
    
    def extract(self, file_path: str) -> DOCXExtractionResult:
        """提取 Word 文档内容"""
        file_path = Path(file_path)
        result = DOCXExtractionResult(file_path=str(file_path))
        
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx 未安装: pip install python-docx")
            return result
        
        try:
            doc = Document(file_path)
            
            # 提取标题
            if doc.core_properties.title:
                result.title = doc.core_properties.title
            
            # 提取段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    result.paragraphs.append(text)
            
            # 提取表格
            for table in doc.tables:
                table_md = self._table_to_markdown(table)
                if table_md:
                    result.tables.append(table_md)
            
            # 合并全文
            result.full_text = "\n\n".join(result.paragraphs)
            
            # 元数据
            core = doc.core_properties
            result.metadata = {
                'author': core.author or "",
                'created': str(core.created) if core.created else "",
                'modified': str(core.modified) if core.modified else "",
                'subject': core.subject or "",
            }
            
            logger.info(f"DOCX 解析完成: {len(result.paragraphs)} 段落, "
                       f"{len(result.tables)} 表格")
            
        except Exception as e:
            logger.error(f"DOCX 解析失败: {e}")
        
        return result
    
    def _table_to_markdown(self, table) -> str:
        """将 Word 表格转换为 Markdown"""
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            row_str = "| " + " | ".join(cells) + " |"
            rows.append(row_str)
            
            if i == 0:
                rows.append("|" + "|".join(["---"] * len(cells)) + "|")
        
        return "\n".join(rows) if rows else ""
